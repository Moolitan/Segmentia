#!/usr/bin/env python3
"""Rewrite paper Sections 3.1/3.2 from the pre-layer-axis DOCX backup."""
from __future__ import annotations

import argparse
import os
import shutil
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


FIGURES = {
    5: "direct_recompute_layerwise_cosine.png",
    6: "direct_recompute_layerwise_normalized_l2.png",
    7: "token_residual_commonality.png",
    8: "corrected_recompute_layerwise_cosine.png",
    9: "corrected_recompute_layerwise_normalized_l2.png",
}


def paragraph_starting_with(document: DocumentType, prefix: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(f"expected one paragraph starting with {prefix!r}, found {len(matches)}")
    return matches[0]


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is None:
        raise ValueError("paragraph is detached")
    parent.remove(paragraph._p)


def remove_range(start: Paragraph, end_exclusive: Paragraph) -> None:
    node = start._p
    while node is not end_exclusive._p:
        next_node = node.getnext()
        if next_node is None:
            raise ValueError("end paragraph is not after start paragraph")
        node.getparent().remove(node)
        node = next_node


def set_run_font(run, size: float | None = None, math: bool = False) -> None:
    font = "Cambria Math" if math else "Times New Roman"
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, run_fonts)
    run_fonts.set(qn("w:ascii"), font)
    run_fonts.set(qn("w:hAnsi"), font)
    run_fonts.set(qn("w:eastAsia"), font if math else "宋体")


def insert_paragraph(
    document: DocumentType,
    reference: Paragraph,
    text: str,
    *,
    style: str = "Normal",
    template: Paragraph | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> Paragraph:
    paragraph = document.add_paragraph(style=style)
    if template is not None and template._p.pPr is not None:
        old_properties = paragraph._p.pPr
        if old_properties is not None:
            paragraph._p.remove(old_properties)
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    set_run_font(run)
    reference._p.addprevious(paragraph._p)
    return paragraph


def insert_heading(
    document: DocumentType,
    reference: Paragraph,
    text: str,
    template: Paragraph,
) -> Paragraph:
    paragraph = insert_paragraph(
        document,
        reference,
        text,
        style="Heading 2",
        template=template,
    )
    for run in paragraph.runs:
        set_run_font(run, size=12)
    return paragraph


def insert_equation(
    document: DocumentType, reference: Paragraph, text: str
) -> Paragraph:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, math=True)
    reference._p.addprevious(paragraph._p)
    return paragraph


def insert_figure(
    document: DocumentType,
    reference: Paragraph,
    image: Path,
    caption: str,
) -> None:
    if not image.is_file():
        raise FileNotFoundError(image)
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(image), width=Inches(6.45))
    reference._p.addprevious(paragraph._p)

    caption_paragraph = document.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = None
    caption_parts = caption.strip().split(maxsplit=2)
    if len(caption_parts) != 3 or caption_parts[0] != "图" or not caption_parts[1].isdigit():
        raise ValueError(f"caption must start with '图 <number>': {caption!r}")

    label_run = caption_paragraph.add_run("图 ")
    begin_run = caption_paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = caption_paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " SEQ 图 \\* ARABIC "
    instruction_run._r.append(instruction)

    separate_run = caption_paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = caption_paragraph.add_run(caption_parts[1])
    end_run = caption_paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    description_run = caption_paragraph.add_run(f" {caption_parts[2]}")
    for run in (
        label_run,
        begin_run,
        instruction_run,
        separate_run,
        result_run,
        end_run,
        description_run,
    ):
        set_run_font(run, size=9)
    reference._p.addprevious(caption_paragraph._p)


def replace_reference_number(
    document: DocumentType,
    paragraph_prefix: str,
    replacements: dict[str, str],
) -> None:
    paragraph = paragraph_starting_with(document, paragraph_prefix)
    for old, new in replacements.items():
        matches = [run for run in paragraph.runs if run.text == old]
        if len(matches) != 1:
            raise ValueError(
                f"expected one run {old!r} in paragraph {paragraph_prefix!r}, "
                f"found {len(matches)}"
            )
        matches[0].text = new


def rewrite(document: DocumentType, figure_dir: Path) -> None:
    heading_32 = paragraph_starting_with(document, "3.2")
    heading_4 = paragraph_starting_with(document, "4 CSKCache设计")
    old_insight = paragraph_starting_with(document, "综上，我们提出如下 insight")
    normal_template = paragraph_starting_with(document, "为进一步理解 Skill KV")
    heading_template = heading_32

    # Replace only the concluding part of 3.1; the Skill-length analysis,
    # behavior table, and attention figures remain untouched from the backup.
    next_after_insight = old_insight._p.getnext()
    remove_paragraph(old_insight)
    if next_after_insight is not None and not Paragraph(next_after_insight, next_after_insight.getparent()).text.strip():
        next_after_insight.getparent().remove(next_after_insight)

    insert_paragraph(
        document,
        heading_32,
        "除注意力路径外，我们进一步从表示层面比较直接复用与完整重计算的差异。对于离线独立 Prefill 得到的 context-free Skill KV，Direct Key 先按照当前请求中的真实插入位置进行 RoPE 对齐，Direct Value 则原样复用；Recompute KV 来自同一 Skill 在真实 OpenHands 请求中完整 Prefill 的结果。图 5 和图 6 分别给出四个代表性 Skill 在 Qwen3-14B 全部 40 层上的余弦相似度与 normalized L2，其中每幅图左侧为 Key，右侧为 Value。",
        template=normal_template,
    )
    insert_figure(
        document,
        heading_32,
        figure_dir / FIGURES[5],
        "图 5  Direct Skill KV 与完整重计算 KV 的逐层余弦相似度",
    )
    insert_paragraph(
        document,
        heading_32,
        "如图 5 所示，即使直接复用跳过了 Skill token 在当前请求中的 Prefill，Key 和 Value 的逐层余弦相似度仍整体保持在 0.9 以上；按 Skill 聚合后，Key 相似度为 0.972–0.984，Value 相似度为 0.960–0.980。图 6 的 normalized L2 进一步表明，两者并非完全一致：Key 误差为 0.149–0.191，Value 误差为 0.187–0.306，并在中后层更为明显。因此，高余弦相似度反映的是缓存表示保留了主要语义方向，而非上下文影响可以忽略。",
        template=normal_template,
    )
    insert_figure(
        document,
        heading_32,
        figure_dir / FIGURES[6],
        "图 6  Direct Skill KV 相对完整重计算 KV 的逐层 normalized L2",
    )
    insert_paragraph(
        document,
        heading_32,
        "上述行为观察、注意力分析和表示相似性共同揭示了 Agent Skill KV 的复用机会。Skill 文本中的流程、工具规范和输出约束具有较强的内部稳定性，离线 KV 已编码大量可被后续 token 直接利用的程序性知识，因此没有必要在每次调用时对完整 Skill 重复 Prefill。另一方面，细粒度规则遗漏、缺失的跨段 Prefill 以及非零 KV 误差说明，Skill 仍需与当前系统角色、用户任务和执行状态对齐。由此，问题从“是否能够复用”转化为“能否用少量连续在线计算恢复必要的请求相关偏差”，这正是下一节关注的核心。",
        template=normal_template,
    )

    # Remove the complete obsolete 3.2 block, including its three layer-axis
    # figures, while preserving Section 4 and everything after it.
    remove_range(heading_32, heading_4)

    insert_heading(
        document,
        heading_4,
        "3.2 基于连续前缀局部观测的 KV 纠错机会",
        heading_template,
    )
    insert_paragraph(
        document,
        heading_4,
        "上一节显示，经过 RoPE 位置对齐的直接复用已经与完整重计算保持较高表示相似度，但仍存在稳定的非零误差。这一现象给出一个比离散 token 选择更直接的启示：若上下文变化只在原有 Skill 表示上引入具有共享结构的残差，则无需重新计算整个 Skill，也无需在长序列中选择分散 token；只要在一个连续局部区间中观测当前请求的残差，就可能估计其余缓存区域所需的修正量。连续区间同时适合 GPU 批量计算，并使在线开销由固定区间长度而非完整 Skill 长度决定。",
        template=normal_template,
    )
    insert_paragraph(
        document,
        heading_4,
        "为检验这一机会，我们以离线 context-free KV 为缓存基线，并在四个真实 OpenHands 任务中采集同一 Skill 的 Recompute KV。设 Xoff 表示离线位置 0 独立 Prefill 得到的缓存，Xdir 表示将缓存 Key 通过 RoPE 旋转到当前插入位置后得到的 Direct KV，Xrec 表示当前请求中完整 Prefill 得到的 KV，其中 X∈{K,V}。对于模型层 l、KV head h 和 Skill 内部 token t，我们定义上下文残差为",
        template=normal_template,
    )
    insert_equation(
        document,
        heading_4,
        "ΔXₜ⁽ˡ˒ʰ⁾ = Xₜ,rec⁽ˡ˒ʰ⁾ − Xₜ,dir⁽ˡ˒ʰ⁾,    X ∈ {K, V}。",
    )
    insert_paragraph(
        document,
        heading_4,
        "实验将 Skill 的前 256 个 token 视为当前请求中连续重计算的校准前缀。为避免起始 wrapper 和 frontmatter 主导估计，我们只使用其后半段 [132,256)；所有指标统一评价从未参与估计的后缀 [256,S)。每个 Skill、每一层、每个 KV head 和 K/V component 均独立估计，不使用其他 Skill、其他模型层或评价后缀的真值。观测区域与后缀区域的平均残差分别为",
        template=normal_template,
    )
    insert_equation(
        document,
        heading_4,
        "μobs⁽ˡ˒ʰ⁾ = (1/124) ∑ₜ₌₁₃₂²⁵⁵ ΔXₜ⁽ˡ˒ʰ⁾,    μsuf⁽ˡ˒ʰ⁾ = Meanₜ∈[256,S) ΔXₜ⁽ˡ˒ʰ⁾。",
    )
    insert_paragraph(
        document,
        heading_4,
        "图 7 以方向余弦衡量观测前缀能否反映未观测后缀的 Key 残差趋势，即比较 μobs 与 μsuf。按 doc-coauthoring、docx、mcp-builder 和 frontend-design 的顺序，逐层、逐 KV head 平均方向余弦分别为 0.793、0.734、0.833 和 0.889，方向为正的单元比例均超过 97.8%。除少量早期层和个别 head 外，热力图呈现大面积深绿色。这说明同一 Skill、同一层、同一 KV head 内，不同 token 的 Key 残差并非相互独立噪声，而包含可由连续局部 token 观测到的公共方向。",
        template=normal_template,
    )
    insert_figure(
        document,
        heading_4,
        figure_dir / FIGURES[7],
        "图 7  同层观测前缀与未观测后缀的 Key 残差方向一致性",
    )
    insert_paragraph(
        document,
        heading_4,
        "直接使用观测均值作为完整修正量会高估部分 Skill 的残差幅值。基于前述诊断，我们在本次实验中统一固定收缩系数 α=0.6，并将同层、同 KV head 的观测偏移应用于未重计算后缀：",
        template=normal_template,
    )
    insert_equation(
        document,
        heading_4,
        "Xₜ,corr⁽ˡ˒ʰ⁾ = Xₜ,dir⁽ˡ˒ʰ⁾ + α μobs⁽ˡ˒ʰ⁾,    t ≥ 256,    α = 0.6。",
    )
    insert_paragraph(
        document,
        heading_4,
        "图 8 和图 9 给出纠错后 KV 与 Recompute KV 的逐层余弦相似度和 normalized L2，并与图 5、图 6 使用相同纵轴范围。对于 Key，四个 Skill 的聚合余弦相似度均提高，绝对增量为 0.0032–0.0128；normalized L2 均降低，降幅为 6.21%–31.45%，因此四个 case 全部同时通过两项表示判据。Value 在三个 case 上改善，但 docx 的余弦相似度下降 0.0004、normalized L2 恶化 0.56%，表明当前公共偏移对 Key 更稳定，尚不能据此声称统一的 K/V 纠错均具有跨 Skill 稳健性。",
        template=normal_template,
    )
    insert_figure(
        document,
        heading_4,
        figure_dir / FIGURES[8],
        "图 8  固定 α=0.6 纠错后 Skill KV 与完整重计算 KV 的逐层余弦相似度",
    )
    insert_figure(
        document,
        heading_4,
        figure_dir / FIGURES[9],
        "图 9  固定 α=0.6 纠错后 Skill KV 相对完整重计算 KV 的逐层 normalized L2",
    )
    insert_paragraph(
        document,
        heading_4,
        "需要强调的是，α=0.6 来自同一批四个 case 的诊断性选择，本节结果用于证明残差方向可共享且统一收缩能够在这些 case 上恢复 Key 表示，不构成对未见 Skill 的泛化结论，也不等价于端到端 Agent 行为或延迟收益。尽管如此，该观察已经给出清晰的系统设计依据：在当前请求中完整计算一个固定长度的连续 Skill 前缀，以其同层、逐 KV head 残差修正剩余缓存 Key，并直接复用未重计算区域。下一章据此设计上下文适配、缓存预取与并发共享机制。",
        template=normal_template,
    )

    # The backup already stores Section 4 captions as automatic Figure 10--13
    # fields, but its prose still contains the pre-insertion Figure 1--5
    # references. Keep all Section 4 content and formatting while restoring the
    # reference numbers used by the immediately preceding paper version.
    replace_reference_number(document, "我们提出了 CSKCache", {" 1 ": " 10 "})
    replace_reference_number(document, "图 1 展示了", {" 1 ": " 10 "})
    replace_reference_number(document, "第 3.2 节的分析表明", {" 2 ": " 11 "})
    replace_reference_number(
        document,
        "如图 2(a) 所示",
        {" 2(a) ": " 11(a) ", " 2(b) ": " 11(b) "},
    )
    replace_reference_number(document, "如图 3(a) 所示", {" 3(a) ": " 12(a) "})
    replace_reference_number(document, "图 3(b) 给出了", {" 3(b) ": " 12(b) "})
    replace_reference_number(document, "为减少并发请求", {" 5 ": " 14 "})
    replace_reference_number(document, "CSKCache 基于 vLLM", {" 5 ": " 14 "})
    replace_reference_number(document, "由于 Skill 复用区间", {" 5 ": " 14 "})


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--base", type=Path, required=True)
    parser.add_argument("--current", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--figure-dir", type=Path, required=True)
    args = parser.parse_args()

    if args.base.resolve() == args.output.resolve():
        raise ValueError("base backup must remain immutable")
    for number, filename in FIGURES.items():
        if not (args.figure_dir / filename).is_file():
            raise FileNotFoundError(f"missing Figure {number}: {args.figure_dir / filename}")

    document = Document(args.base)
    rewrite(document, args.figure_dir)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    backup_path: Path | None = None
    if args.current.is_file() and args.current.resolve() == args.output.resolve():
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        backup_path = args.current.with_name(
            f"{args.current.stem}.backup-before-token-axis-rewrite-{timestamp}.docx"
        )
        shutil.copy2(args.current, backup_path)

    temporary = args.output.with_name(f".{args.output.name}.tmp")
    document.save(temporary)
    # A second parse catches broken relationships or malformed document XML
    # before the requested output is replaced.
    reparsed = Document(temporary)
    expected_captions = {f"图 {number}" for number in FIGURES}
    actual_captions = {
        p.text.strip().split(maxsplit=2)[0] + " " + p.text.strip().split(maxsplit=2)[1]
        for p in reparsed.paragraphs
        if p.text.strip().startswith("图 ") and len(p.text.strip().split()) >= 2
    }
    if not expected_captions.issubset(actual_captions):
        raise ValueError(f"missing captions: {sorted(expected_captions - actual_captions)}")
    os.replace(temporary, args.output)
    print(f"[paper-rewritten] base={args.base}")
    print(f"[paper-rewritten] backup={backup_path}")
    print(f"[paper-rewritten] output={args.output}")
    print(
        f"[paper-rewritten] paragraphs={len(reparsed.paragraphs)} "
        f"tables={len(reparsed.tables)} images={len(reparsed.inline_shapes)}"
    )


if __name__ == "__main__":
    main()
